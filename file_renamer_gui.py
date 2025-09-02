#!/usr/bin/env python3
"""
智能文件重命名桌面应用
支持图片OCR、PDF/文档内容识别、批量重命名、回滚功能
"""

import sys
import os
import re
from pathlib import Path
from typing import List, Optional, Dict, Any
import json
import shutil
from datetime import datetime

# PyQt6 imports
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QTextEdit, QProgressBar, QFileDialog,
    QMessageBox, QGroupBox, QCheckBox, QSpinBox, QComboBox,
    QLineEdit, QTabWidget, QTableWidget, QTableWidgetItem,
    QHeaderView, QSplitter, QFrame
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt6.QtGui import QFont, QIcon, QPixmap

# 重命名逻辑导入
try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
    from docx import Document
except ImportError:
    docx = None
    Document = None

try:
    import chardet
except ImportError:
    chardet = None

from tqdm import tqdm


class FileRenamer:
    """文件重命名核心逻辑类"""
    
    def __init__(self, log_callback=None):
        self.image_exts = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
        self.pdf_exts = {".pdf"}
        self.docx_exts = {".docx"}
        self.txt_exts = {".txt", ".md", ".csv"}
        self.log_callback = log_callback  # 添加日志回调函数
    
    def _log(self, message):
        """统一的日志输出方法"""
        if self.log_callback:
            self.log_callback(message)
        print(f"[FileRenamer] {message}")
    
    def detect_file_type(self, path: Path) -> str:
        """检测文件类型"""
        suffix = path.suffix.lower()
        if suffix in self.image_exts:
            return "image"
        elif suffix in self.pdf_exts:
            return "pdf"
        elif suffix in self.docx_exts:
            return "docx"
        elif suffix in self.txt_exts:
            return "txt"
        return "other"
    
    def read_image_text(self, path: Path, extract_len: int) -> str:
        """读取图片文本（OCR）"""
        if pytesseract is None or Image is None:
            return ""
        try:
            with Image.open(str(path)) as img:
                text = pytesseract.image_to_string(img)
            return (text or "")[:extract_len]
        except Exception:
            return ""
    
    def read_pdf_text(self, path: Path, pages: int, extract_len: int) -> str:
        """读取PDF文本"""
        if PdfReader is None:
            return ""
        try:
            reader = PdfReader(str(path))
            acc = []
            for i, page in enumerate(reader.pages[:pages]):
                try:
                    acc.append(page.extract_text() or "")
                except Exception:
                    continue
            text = " ".join(acc)
            return text[:extract_len]
        except Exception:
            return ""
    
    def read_docx_text(self, path: Path, extract_len: int) -> str:
        """读取DOCX文本"""
        if docx is None:
            return ""
        try:
            d = docx.Document(str(path))
            paragraphs = [p.text for p in d.paragraphs if p.text]
            text = " ".join(paragraphs)
            return text[:extract_len]
        except Exception:
            return ""
    
    def read_txt_text(self, path: Path, extract_len: int) -> str:
        """读取TXT文本"""
        if chardet is not None:
            try:
                with open(path, "rb") as f:
                    raw = f.read(4096)
                enc = chardet.detect(raw).get("encoding") or "utf-8"
            except Exception:
                enc = "utf-8"
        else:
            enc = "utf-8"
        try:
            with open(path, "r", encoding=enc, errors="ignore") as f:
                text = f.read(extract_len)
            return text
        except Exception:
            return ""
    
    def normalize_name(self, text: str, lowercase: bool = True, 
                      replace_space_with_underscore: bool = True, 
                      max_length: int = 60) -> str:
        """规范化文件名"""
        if not text:
            return "unnamed"
        
        import re
        # 清理空白字符
        text = re.sub(r'\s+', ' ', text).strip()
        if replace_space_with_underscore:
            text = text.replace(" ", "_")
        
        # 移除非法字符
        invalid_chars = r'<>:"/\|?*'
        text = re.sub(f'[{re.escape(invalid_chars)}]', '', text)
        
        # 清理分隔符
        text = re.sub(r'[_\-\s]+', '_', text)
        text = text.rstrip(" .")
        
        if lowercase:
            text = text.lower()
        
        if not text:
            text = "unnamed"
        
        if len(text) > max_length:
            text = text[:max_length]
        
        return text
    
    def propose_new_name(self, path: Path, extract_len: int = 120,
                        lowercase: bool = True, space_to_underscore: bool = True,
                        max_length: int = 60) -> str:
        """生成新文件名，优先调用 DeepSeek；失败时退回启发式。"""
        # 1) 统一调用 DeepSeek
        try:
            from deepseek_api_service import deepseek_service
            deepseek_result = None
            if deepseek_service.is_available():
                deepseek_result = deepseek_service.extract_renaming_info(path)
                if deepseek_result:
                    return deepseek_result
                # 记录失败原因供UI显示
                if getattr(deepseek_service, 'last_error', None):
                    self._log(f"DeepSeek失败: {deepseek_service.last_error}")
                if getattr(deepseek_service, 'last_suggestion', None):
                    self._log(f"建议: {deepseek_service.last_suggestion}")
        except Exception as e:
            self._log(f"DeepSeek调用异常: {e}")
        
        # 2) 退回到原有金融专用提取逻辑
        base_text = self.extract_content_for_naming(path)
        if not base_text:
            base_text = path.stem
        safe_stem = self.normalize_name(
            base_text,
            lowercase=lowercase,
            replace_space_with_underscore=space_to_underscore,
            max_length=max_length
        )
        return f"{safe_stem}{path.suffix.lower()}"
    
    def make_unique_path(self, target_path: Path) -> Path:
        """生成唯一路径（避免重名）"""
        if not target_path.exists():
            return target_path
        
        stem = target_path.stem
        suffix = target_path.suffix
        parent = target_path.parent
        counter = 1
        
        while True:
            candidate = parent / f"{stem}-{counter}{suffix}"
            if not candidate.exists():
                return candidate
            counter += 1

    def extract_content_for_naming(self, file_path):
        """提取文件内容用于重命名，专门针对金融文档优化"""
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            # 针对金融文档的智能内容提取
            if file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                return self._extract_image_content_financial(file_path)
            elif file_extension == '.pdf':
                return self._extract_pdf_content_financial(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_content_financial(file_path)
            elif file_extension in ['.txt', '.csv', '.xlsx', '.xls']:
                return self._extract_text_content_financial(file_path)
            else:
                # 其他文件类型，尝试提取文件名中的关键信息
                return self._extract_filename_info_financial(file_path)
                
        except Exception as e:
            self._log(f"内容提取失败 {file_path}: {str(e)}")
            return file_path.stem
    
    def _extract_image_content_financial(self, file_path):
        """从图片中提取金融相关信息，优先使用DeepSeek API"""
        try:
            # 1. 优先尝试DeepSeek API
            try:
                from deepseek_api_service import deepseek_service
                if deepseek_service.is_available():
                    deepseek_result = deepseek_service.extract_renaming_info(file_path)
                    if deepseek_result:
                        self._log(f"DeepSeek API识别成功: {deepseek_result}")
                        return deepseek_result.replace(file_path.suffix, '')  # 返回不带扩展名的名称
            except ImportError:
                pass  # DeepSeek服务未安装
            except Exception as e:
                self._log(f"DeepSeek API调用失败，回退到传统OCR: {e}")
                pass  # DeepSeek失败，继续使用传统OCR
            
            # 2. 备选方案：AI OCR服务（已禁用，避免下载模型）
            # 注释掉AI OCR服务，避免下载耗时模型
            # try:
            #     from ai_ocr_service import ai_ocr_service
            #     ai_text = ai_ocr_service.extract_text_from_image(file_path)
            #     if ai_text and len(ai_text.strip()) > 20:
            #         extracted_info = self._extract_financial_keywords(ai_text)
            #         if extracted_info and len(extracted_info.strip()) > 10:
            #             return extracted_info
            # except ImportError:
            #     pass  # AI OCR服务未安装
            # except Exception as e:
            #     self._log(f"AI OCR服务调用失败，回退到传统OCR: {e}")
            #     pass  # AI OCR失败，继续使用传统OCR
            
            # 2. 使用传统OCR（pytesseract）
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            
            # 针对金融文档的关键信息提取
            extracted_info = self._extract_financial_keywords(text)
            if extracted_info and len(extracted_info.strip()) > 10:  # 确保提取的信息有意义
                return extracted_info
            
            # 3. 如果OCR结果不够好，优先使用文件名信息
            filename_info = self._extract_filename_info_financial(file_path)
            if filename_info and filename_info != file_path.stem:
                return filename_info
            
            # 4. 最后才使用OCR结果
            return extracted_info if extracted_info else file_path.stem
            
        except Exception as e:
            self._log(f"图片OCR失败 {file_path}: {str(e)}")
            return self._extract_filename_info_financial(file_path)
    
    def _extract_pdf_content_financial(self, file_path):
        """从PDF中提取金融相关信息，优先使用DeepSeek API"""
        try:
            # 1. 优先尝试DeepSeek API
            try:
                from deepseek_api_service import deepseek_service
                if deepseek_service.is_available():
                    deepseek_result = deepseek_service.extract_renaming_info(file_path)
                    if deepseek_result:
                        self._log(f"DeepSeek API识别成功: {deepseek_result}")
                        return deepseek_result.replace(file_path.suffix, '')  # 返回不带扩展名的名称
            except ImportError:
                pass  # DeepSeek服务未安装
            except Exception as e:
                self._log(f"DeepSeek API调用失败，回退到传统方法: {e}")
                pass  # DeepSeek失败，继续使用传统方法
            
            # 2. 备选方案：AI OCR服务（处理扫描版PDF，已禁用，避免下载模型）
            # 注释掉AI OCR服务，避免下载耗时模型
            # try:
            #     from ai_ocr_service import ai_ocr_service
            #     ai_text = ai_ocr_service.extract_text_from_pdf(file_path)
            #     if ai_text and len(ai_text.strip()) > 20:
            #         extracted_info = self._extract_financial_keywords(ai_text)
            #         if extracted_info and len(extracted_info.strip()) > 10:
            #             return extracted_info
            # except ImportError:
            #     pass  # AI OCR服务未安装
            # except Exception as e:
            #     self._log(f"PDF AI OCR服务调用失败，回退到传统方法: {e}")
            #     pass  # AI OCR失败，继续使用传统方法
            
            # 2. 使用pdfplumber（更好的文本提取）
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page_num in range(min(3, len(pdf.pages))):
                        page = pdf.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    
                    # 如果pdfplumber提取到文本，使用它
                    if text.strip():
                        extracted_info = self._extract_financial_keywords(text)
                        if extracted_info and len(extracted_info.strip()) > 10:
                            return extracted_info
            except ImportError:
                pass  # pdfplumber未安装，继续使用pypdf
            except Exception as e:
                self._log(f"pdfplumber提取失败 {file_path}: {str(e)}")
            
            # 3. 使用pypdf作为备选方案
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text = ""
                
                # 提取前几页内容（通常关键信息在前几页）
                for page_num in range(min(3, len(pdf_reader.pages))):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                # 针对金融文档的关键信息提取
                extracted_info = self._extract_financial_keywords(text)
                if extracted_info and len(extracted_info.strip()) > 10:  # 确保提取的信息有意义
                    return extracted_info
            
            # 4. 如果文本提取都失败了，尝试从文件名提取更多信息
            filename_info = self._extract_filename_info_financial(file_path)
            if filename_info and filename_info != file_path.stem:
                return filename_info
            
            # 5. 最后才使用提取的文本
            return extracted_info if extracted_info else file_path.stem
                
        except Exception as e:
            self._log(f"PDF解析失败 {file_path}: {str(e)}")
            return self._extract_filename_info_financial(file_path)
    
    def _extract_docx_content_financial(self, file_path):
        """从DOCX中提取金融相关信息"""
        try:
            doc = Document(file_path)
            text = ""
            
            # 提取文档内容
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 针对金融文档的关键信息提取
            extracted_info = self._extract_financial_keywords(text)
            if extracted_info:
                return extracted_info
            
            # 如果内容提取失败，尝试从文件名提取
            return self._extract_filename_info_financial(file_path)
            
        except Exception as e:
            self._log(f"DOCX解析失败 {file_path}: {str(e)}")
            return self._extract_filename_info_financial(file_path)
    
    def _extract_text_content_financial(self, file_path):
        """从文本文件中提取金融相关信息"""
        try:
            # 检测文件编码
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected_encoding = chardet.detect(raw_data)['encoding']
            
            # 尝试不同编码读取文件
            encodings_to_try = [detected_encoding, 'utf-8', 'gbk', 'gb2312', 'latin-1']
            text = ""
            
            for encoding in encodings_to_try:
                if encoding:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                            break
                    except UnicodeDecodeError:
                        continue
            
            # 针对金融文档的关键信息提取
            extracted_info = self._extract_financial_keywords(text)
            if extracted_info:
                return extracted_info
            
            # 如果内容提取失败，尝试从文件名提取
            return self._extract_filename_info_financial(file_path)
            
        except Exception as e:
            self._log(f"文本文件读取失败 {file_path}: {str(e)}")
            return self._extract_filename_info_financial(file_path)
    
    def _extract_financial_keywords(self, text):
        """从文本中提取金融文档的关键信息"""
        if not text:
            return None
        
        # 清理文本
        text = text.strip()
        if len(text) < 10:  # 文本太短，无法提取有效信息
            return None
        
        # 1. 提取基金名称（更智能的识别）
        fund_patterns = [
            # 完整的基金名称模式
            r'([^，。\n]{2,40}(?:私募证券投资基金|私募基金|证券投资基金|基金))',
            # 包含期数的基金名称（改进版）
            r'([^，。\n]{2,40}(?:1号|2号|3号|4号|5号|6号|7号|8号|9号)[^，。\n]{0,20}(?:期)?(?:私募证券投资基金|私募基金|证券投资基金|基金))',
            # 包含策略的基金名称（改进版）
            r'([^，。\n]{2,40}(?:多策略|稳进|进取|平衡|稳健|成长|价值|量化|CTA|FOF)[^，。\n]{0,20}(?:私募证券投资基金|私募基金|证券投资基金|基金))',
            # 简化基金名称（新增，匹配"展弘稳进1号7期私募基金"）
            r'([^，。\n]{2,20}(?:稳进|多策略|进取|平衡|稳健|成长|价值|量化|CTA|FOF)[^，。\n]{0,15}(?:1号|2号|3号|4号|5号|6号|7号|8号|9号)[^，。\n]{0,15}(?:期)?[^，。\n]{0,20}(?:私募基金|基金))',
            # 投资管理公司
            r'([^，。\n]{2,30}(?:投资管理|资产管理|基金管理)(?:有限公司|股份公司|有限责任公司))',
            # 一般公司名称
            r'([^，。\n]{2,30}(?:有限公司|股份公司|有限责任公司))'
        ]
        
        fund_name = None
        for pattern in fund_patterns:
            matches = re.findall(pattern, text)
            if matches:
                fund_name = matches[0].strip()
                break
        
        # 2. 提取客户姓名（通常是2-4个中文字符）
        name_patterns = [
            r'姓名[：:]\s*([一-龯]{2,4})',
            r'客户[：:]\s*([一-龯]{2,4})',
            r'([一-龯]{2,4})\s*先生',
            r'([一-龯]{2,4})\s*女士',
            r'([一-龯]{2,4})',  # 2-4个中文字符（放在最后，避免误匹配）
        ]
        
        client_name = None
        for pattern in name_patterns:
            matches = re.findall(pattern, text)
            if matches:
                # 过滤掉常见的非姓名词汇
                candidate = matches[0].strip()
                if not any(word in candidate for word in ['基金', '投资', '管理', '公司', '有限', '股份', '私募', '证券', '策略', '号']):
                    client_name = candidate
                    break
        
        # 3. 提取文件类型/用途（更智能的识别）
        doc_type_patterns = [
            # 业务凭证类
            r'(打款凭证|转账凭证|汇款凭证|付款凭证|收款凭证|银行回单)',
            # 客户资料类
            r'(基本信息表|客户信息表|个人信息表|资料表|客户资料|个人资料)',
            # 法律文件类
            r'(合同|协议|委托书|确认函|确认书|授权书|承诺函)',
            # 身份证明类
            r'(身份证|护照|户口本|结婚证|证件|身份证明)',
            # 财务记录类
            r'(银行流水|对账单|存单|存折|流水|财务记录)',
            # 风险文件类
            r'(风险提示|风险告知|风险确认|风险书|风险揭示书)',
            # 投资操作类
            r'(认购|申购|赎回|转换|投资|交易)',
            # 收益相关类
            r'(收益分配|分红|派息|收益|收益确认)',
            # 申请表格类
            r'(申请表|登记表|备案表|审核表|审批表)',
            # 通知说明类
            r'(通知书|告知书|说明|报告|公告|通知)',
            # 基金相关类
            r'(基金合同|基金招募说明书|基金说明书|基金公告|基金报告)',
            # 基金公告类（新增，更精确的匹配）
            r'(临时开放日公告|开放日公告|定期开放公告|申购赎回公告|分红公告|净值公告|收益公告|风险提示公告|投资策略公告|基金经理变更公告|基金公告)',
            # 公告类（通用，放在基金公告类之后）
            r'(公告|通知|通告|公示|声明)',
            # 其他业务类
            r'(业务确认|业务回执|业务凭证|业务单据)'
        ]
        
        doc_type = None
        for pattern in doc_type_patterns:
            matches = re.findall(pattern, text)
            if matches:
                doc_type = matches[0].strip()
                break
        
        # 4. 提取日期信息
        date_patterns = [
            r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',  # YYYY-MM-DD 或 YYYY/M/D
            r'(\d{4}年\d{1,2}月\d{1,2}日)',      # YYYY年MM月DD日
            r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})'   # MM/DD/YYYY
        ]
        
        date_info = None
        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            if matches:
                date_info = matches[0].strip()
                break
        
        # 5. 组合关键信息（优化版，避免重复）
        parts = []
        if fund_name:
            parts.append(fund_name)
        if client_name:
            # 避免客户姓名与基金名称重复
            if not fund_name or client_name not in fund_name:
                parts.append(client_name)
        if doc_type:
            # 避免文档类型与基金名称重复
            if not fund_name or doc_type not in fund_name:
                parts.append(doc_type)
        if date_info:
            parts.append(date_info)
        
        if parts:
            # 限制总长度，避免文件名过长
            result = "-".join(parts)
            if len(result) > 100:
                result = result[:100]
            return result
        
        # 如果没有提取到关键信息，返回前50个字符
        return text[:50].strip()
    
    def _extract_filename_info_financial(self, file_path):
        """从文件名中提取金融相关信息"""
        filename = file_path.stem
        
        # 如果文件名已经包含关键信息，直接使用
        if any(keyword in filename for keyword in ['私募', '基金', '投资', '管理', '微信', '图片']):
            return filename
        
        # 尝试从文件名中提取信息
        # 1. 提取日期（支持多种格式）
        date_patterns = [
            r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',  # YYYY-MM-DD 或 YYYY/M/D
            r'(\d{8})',                           # YYYYMMDD
            r'(\d{4}年\d{1,2}月\d{1,2}日)',       # YYYY年MM月DD日
            r'(\d{4})(\d{2})(\d{2})',            # YYYYMMDD（分别捕获年月日）
            r'(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})',  # YYYYMMDDHHMMSS
        ]
        
        for pattern in date_patterns:
            date_match = re.search(pattern, filename)
            if date_match:
                if len(date_match.groups()) == 1:
                    # 单个日期字符串
                    date_str = date_match.group(1)
                    if len(date_str) == 8:  # YYYYMMDD格式
                        formatted_date = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
                    else:
                        formatted_date = date_str
                    return f"金融文档-{formatted_date}"
                elif len(date_match.groups()) == 3:
                    # YYYYMMDD分别捕获
                    year, month, day = date_match.groups()
                    formatted_date = f"{year}-{month}-{day}"
                    return f"金融文档-{formatted_date}"
                elif len(date_match.groups()) == 6:
                    # YYYYMMDDHHMMSS分别捕获
                    year, month, day, hour, minute, second = date_match.groups()
                    formatted_date = f"{year}-{month}-{day}_{hour}-{minute}-{second}"
                    return f"金融文档-{formatted_date}"
        
        # 2. 提取可能的客户姓名（2-4个中文字符）
        name_match = re.search(r'([一-龯]{2,4})', filename)
        if name_match:
            return f"金融文档-{name_match.group(1)}"
        
        # 3. 如果都没有，返回原文件名
        return filename
    
    def log_message(self, message: str):
        """记录日志消息（GUI模式下输出到控制台）"""
        self._log(f"[FileRenamer] {message}")


class RenameWorker(QThread):
    """重命名工作线程"""
    progress_updated = pyqtSignal(int, int)  # current, total
    file_processed = pyqtSignal(str, str)    # old_name, new_name
    finished = pyqtSignal(dict)              # results
    error_occurred = pyqtSignal(str)         # error message
    
    def __init__(self, source_paths: List[Path], target_dir: Path, 
                 config: Dict[str, Any], copy_mode: bool = False, log_callback=None):
        super().__init__()
        self.source_paths = source_paths
        self.target_dir = target_dir
        self.config = config
        self.copy_mode = copy_mode
        self.log_callback = log_callback
        self.renamer = FileRenamer(log_callback=log_callback)
        self.rename_log = []
    
    def log_message(self, message: str):
        """记录日志消息（GUI模式下输出到控制台和执行日志区域）"""
        if self.log_callback:
            self.log_callback(message)
        print(f"[RenameWorker] {message}")
    
    def run(self):
        try:
            results = self._process_files()
            self.finished.emit(results)
        except Exception as e:
            self.error_occurred.emit(str(e))
    
    def _process_files(self) -> Dict[str, Any]:
        """处理文件重命名"""
        all_files = []
        
        # 收集所有文件
        for source_path in self.source_paths:
            if source_path.is_file():
                all_files.append(source_path)
            elif source_path.is_dir():
                for file_path in source_path.rglob("*"):
                    if file_path.is_file():
                        all_files.append(file_path)
        
        # 过滤文件类型
        if self.config.get('include_exts'):
            include_set = {ext.lower() if ext.startswith('.') else f'.{ext.lower()}' 
                          for ext in self.config['include_exts']}
            all_files = [f for f in all_files if f.suffix.lower() in include_set]
        
        total_files = len(all_files)
        processed = 0
        success_count = 0
        error_count = 0
        
        for file_path in all_files:
            try:
                # 生成新文件名
                new_name = self.renamer.propose_new_name(
                    file_path,
                    extract_len=self.config.get('extract_len', 120),
                    lowercase=self.config.get('lowercase', True),
                    space_to_underscore=self.config.get('space_to_underscore', True),
                    max_length=self.config.get('max_length', 60)
                )
                
                # 确定目标路径
                if self.copy_mode:
                    target_path = self.target_dir / new_name
                    target_path = self.renamer.make_unique_path(target_path)
                    
                    # 复制文件
                    shutil.copy2(file_path, target_path)
                    action = "copied"
                else:
                    target_path = file_path.with_name(new_name)
                    target_path = self.renamer.make_unique_path(target_path)
                    
                    # 移动文件
                    file_path.rename(target_path)
                    action = "renamed"
                
                # 记录操作
                self.rename_log.append({
                    'old_path': str(file_path),
                    'new_path': str(target_path),
                    'action': action,
                    'timestamp': datetime.now().isoformat()
                })
                
                # 发送进度信号
                processed += 1
                self.progress_updated.emit(processed, total_files)
                self.file_processed.emit(file_path.name, target_path.name)
                
                success_count += 1
                
            except Exception as e:
                error_count += 1
                self.log_message(f"处理文件 {file_path.name} 时出错: {str(e)}")
        
        return {
            'total': total_files,
            'success': success_count,
            'error': error_count,
            'rename_log': self.rename_log
        }


class FileRenamerGUI(QMainWindow):
    """文件重命名主界面"""
    
    def __init__(self):
        super().__init__()
        self.source_paths = []
        self.target_dir = None
        self.rename_log = []
        self.worker = None
        
        self.init_ui()
        self.load_config()
        
        # 设置 DeepSeek API 服务的日志回调
        try:
            from deepseek_api_service import deepseek_service
            deepseek_service.set_log_callback(self.log_message)
        except ImportError:
            pass  # DeepSeek 服务未安装
    
    def log_message(self, message: str):
        """记录日志消息（GUI模式下输出到控制台和执行日志区域）"""
        print(f"[FileRenamerGUI] {message}")
        # 如果GUI已初始化，也输出到执行日志区域
        if hasattr(self, 'log_text') and self.log_text:
            self.log_text.append(f"[{datetime.now().strftime('%H:%M:%S')}] {message}")
            # 滚动到底部
            self.log_text.verticalScrollBar().setValue(
                self.log_text.verticalScrollBar().maximum()
            )
    
    def init_ui(self):
        """初始化界面"""
        self.setWindowTitle("智能文件重命名工具")
        self.setGeometry(100, 100, 900, 400)
        
        # 设置窗口样式
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f8f9fa;
            }
            QTabWidget::pane {
                border: 1px solid #e0e0e0;
                background-color: white;
                border-radius: 8px;
            }
            QTabBar::tab {
                background-color: #ecf0f1;
                color: #2c3e50;
                padding: 12px 20px;
                margin-right: 2px;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
                font-weight: bold;
            }
            QTabBar::tab:selected {
                background-color: #3498db;
                color: white;
            }
            QTabBar::tab:hover {
                background-color: #bdc3c7;
            }
        """)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QVBoxLayout(central_widget)
        
        # 创建标签页
        tab_widget = QTabWidget()
        main_layout.addWidget(tab_widget)
        
        # 主操作标签页
        main_tab = self.create_main_tab()
        tab_widget.addTab(main_tab, "重命名操作")
        
        # 配置标签页
        config_tab = self.create_config_tab()
        tab_widget.addTab(config_tab, "配置选项")
        
        # 结果标签页
        result_tab = self.create_result_tab()
        tab_widget.addTab(result_tab, "执行结果")
        
        # 状态栏
        # self.statusBar().showMessage("就绪")  # 移除状态栏显示
    
    def create_main_tab(self) -> QWidget:
        """创建主操作标签页"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(12)
        layout.setContentsMargins(12, 12, 12, 12)
        
        # 文件选择区域
        file_group = QGroupBox("📁 选择要重命名的文件/文件夹")
        file_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
        """)
        file_layout = QVBoxLayout(file_group)
        file_layout.setSpacing(10)
        
        # 统一文件选择按钮
        file_btn_layout = QHBoxLayout()
        self.select_files_btn = QPushButton("�� 选择文件或文件夹")
        self.select_files_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:pressed {
                background-color: #21618c;
            }
        """)
        self.select_files_btn.clicked.connect(self.select_files_or_folders)
        
        self.clear_selection_btn = QPushButton("🗑️ 清除选择")
        self.clear_selection_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
            QPushButton:pressed {
                background-color: #a93226;
            }
        """)
        self.clear_selection_btn.clicked.connect(self.clear_selection)
        
        file_btn_layout.addWidget(self.select_files_btn)
        file_btn_layout.addWidget(self.clear_selection_btn)
        file_btn_layout.addStretch()
        file_layout.addLayout(file_btn_layout)
        
        # 已选择的文件列表
        self.file_list_label = QLabel("未选择任何文件")
        self.file_list_label.setStyleSheet("""
            QLabel {
                color: #7f8c8d;
                font-size: 13px;
                padding: 10px;
                background-color: #f8f9fa;
                border-radius: 4px;
                border: 1px solid #e9ecef;
            }
        """)
        file_layout.addWidget(self.file_list_label)
        
        layout.addWidget(file_group)
        
        # 目标目录选择
        target_group = QGroupBox("📂 选择目标目录")
        target_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
        """)
        target_layout = QHBoxLayout(target_group)
        target_layout.setSpacing(10)
        
        self.target_path_edit = QLineEdit()
        self.target_path_edit.setPlaceholderText("选择重命名后文件的存储目录")
        self.target_path_edit.setReadOnly(True)
        self.target_path_edit.setStyleSheet("""
            QLineEdit {
                padding: 10px;
                border: 2px solid #e0e0e0;
                border-radius: 6px;
                font-size: 13px;
                background-color: #f8f9fa;
            }
            QLineEdit:focus {
                border-color: #3498db;
            }
        """)
        
        self.select_target_btn = QPushButton("📁 选择目录")
        self.select_target_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:pressed {
                background-color: #1e8449;
            }
        """)
        self.select_target_btn.clicked.connect(self.select_target_directory)
        
        target_layout.addWidget(self.target_path_edit)
        target_layout.addWidget(self.select_target_btn)
        layout.addWidget(target_group)
        
        # 存储模式选择
        mode_group = QGroupBox("💾 存储模式")
        mode_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
        """)
        mode_layout = QVBoxLayout(mode_group)
        mode_layout.setSpacing(10)
        
        # 使用单选按钮
        from PyQt6.QtWidgets import QRadioButton
        
        self.copy_mode_radio = QRadioButton("📋 复制模式（保留原文件，重命名后存储到目标目录）")
        self.copy_mode_radio.setChecked(True)  # 默认选中复制模式
        self.copy_mode_radio.setStyleSheet("""
            QRadioButton {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QRadioButton::indicator {
                width: 20px;
                height: 20px;
            }
            QRadioButton::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 10px;
                background-color: white;
            }
            QRadioButton::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 10px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iOCIgdmlld0JveD0iMCAwIDggOCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iNCIgY3k9IjQiIHI9IjIiIGZpbGw9IndoaXRlIi8+Cjwvc3ZnPgo=);
            }
        """)
        
        self.overwrite_mode_radio = QRadioButton("✏️ 覆盖模式（直接重命名原文件）")
        self.overwrite_mode_radio.setStyleSheet("""
            QRadioButton {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QRadioButton::indicator {
                width: 20px;
                height: 20px;
            }
            QRadioButton::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 10px;
                background-color: white;
            }
            QRadioButton::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 10px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iOCIgdmlld0JveD0iMCAwIDggOCIgZmlsbD0ibm9uZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIj4KPGNpcmNsZSBjeD0iNCIgY3k9IjQiIHI9IjIiIGZpbGw9IndoaXRlIi8+Cjwvc3ZnPgo=);
            }
        """)
        
        mode_layout.addWidget(self.copy_mode_radio)
        mode_layout.addWidget(self.overwrite_mode_radio)
        layout.addWidget(mode_group)
        
        # 执行按钮
        execute_layout = QHBoxLayout()
        self.execute_btn = QPushButton("开始重命名")
        self.execute_btn.clicked.connect(self.start_rename)
        self.execute_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-size: 14px; padding: 10px; }")
        
        self.rollback_btn = QPushButton("回滚操作")
        self.rollback_btn.clicked.connect(self.rollback_operations)
        self.rollback_btn.setEnabled(False)
        
        execute_layout.addWidget(self.execute_btn)
        execute_layout.addWidget(self.rollback_btn)
        execute_layout.addStretch()
        layout.addLayout(execute_layout)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # 实时日志
        log_group = QGroupBox("执行日志")
        log_layout = QVBoxLayout(log_group)
        self.log_text = QTextEdit()
        self.log_text.setMaximumHeight(80)
        log_layout.addWidget(self.log_text)
        layout.addWidget(log_group)
        
        layout.addStretch()
        return widget
    
    def create_config_tab(self) -> QWidget:
        """创建配置标签页"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(12)
        layout.setContentsMargins(12, 12, 12, 12)
        
        # DeepSeek API配置
        api_group = QGroupBox("🤖 DeepSeek API配置")
        api_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
        """)
        api_layout = QVBoxLayout(api_group)
        api_layout.setSpacing(10)
        
        # API密钥输入
        api_key_layout = QHBoxLayout()
        api_key_label = QLabel("API密钥:")
        api_key_label.setStyleSheet("font-weight: bold; color: #2c3e50; font-size: 13px;")
        api_key_layout.addWidget(api_key_label)
        
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("请输入你的DeepSeek API密钥")
        self.api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.api_key_input.setStyleSheet("""
            QLineEdit {
                padding: 10px;
                border: 2px solid #e0e0e0;
                border-radius: 6px;
                font-size: 13px;
                background-color: #f8f9fa;
            }
            QLineEdit:focus {
                border-color: #3498db;
            }
        """)
        
        # 显示/隐藏明文按钮
        self.toggle_password_btn = QPushButton("👁")
        self.toggle_password_btn.setMaximumWidth(35)
        self.toggle_password_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                border-radius: 4px;
                font-size: 12px;
                padding: 8px;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        self.toggle_password_btn.setToolTip("显示/隐藏API密钥")
        self.toggle_password_btn.clicked.connect(self.toggle_password_visibility)
        
        api_key_layout.addWidget(self.api_key_input)
        api_key_layout.addWidget(self.toggle_password_btn)
        api_layout.addLayout(api_key_layout)
        
        # API配置按钮
        api_btn_layout = QHBoxLayout()
        self.test_api_btn = QPushButton("🧪 测试API")
        self.test_api_btn.setStyleSheet("""
            QPushButton {
                background-color: #f39c12;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #e67e22;
            }
            QPushButton:pressed {
                background-color: #d35400;
            }
        """)
        self.test_api_btn.clicked.connect(self.test_api_key)
        
        self.clear_api_btn = QPushButton("🗑️ 清空")
        self.clear_api_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
            QPushButton:pressed {
                background-color: #a93226;
            }
        """)
        self.clear_api_btn.clicked.connect(self.clear_api_key)
        
        api_btn_layout.addWidget(self.test_api_btn)
        api_btn_layout.addWidget(self.clear_api_btn)
        api_layout.addLayout(api_btn_layout)
        
        # API状态显示
        self.api_status_label = QLabel("API状态: 未配置")
        self.api_status_label.setStyleSheet("""
            QLabel {
                color: #7f8c8d;
                font-size: 13px;
                padding: 8px;
                background-color: #ecf0f1;
                border-radius: 4px;
                border: 1px solid #bdc3c7;
            }
        """)
        api_layout.addWidget(self.api_status_label)
        
        # API帮助信息
        help_text = QTextEdit()
        help_text.setMaximumHeight(100)
        help_text.setPlainText("""
如何获取API密钥：
1. 访问 https://platform.deepseek.com/
2. 注册并登录账户
3. 进入API管理页面
4. 创建新的API密钥
5. 复制密钥并粘贴到上面的输入框

注意：请妥善保管你的API密钥，不要分享给他人。
        """)
        help_text.setReadOnly(True)
        help_text.setStyleSheet("""
            QTextEdit {
                background-color: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 6px;
                padding: 10px;
                font-size: 12px;
                color: #495057;
            }
        """)
        api_layout.addWidget(help_text)
        
        layout.addWidget(api_group)
        
        # 创建水平布局容器，将重命名规则和文件类型过滤并排
        rules_filter_layout = QHBoxLayout()
        rules_filter_layout.setSpacing(15)
        
        # 重命名规则配置
        rules_group = QGroupBox("⚙️ 重命名规则")
        rules_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
        """)
        rules_layout = QVBoxLayout(rules_group)
        rules_layout.setSpacing(10)
        
        # 文本提取长度
        extract_len_layout = QHBoxLayout()
        extract_len_label = QLabel("文本提取长度:")
        extract_len_label.setStyleSheet("font-weight: bold; color: #2c3e50; font-size: 13px;")
        extract_len_layout.addWidget(extract_len_label)
        
        self.extract_len_spin = QSpinBox()
        self.extract_len_spin.setRange(50, 500)
        self.extract_len_spin.setValue(120)
        self.extract_len_spin.setStyleSheet("""
            QSpinBox {
                padding: 8px;
                border: 2px solid #e0e0e0;
                border-radius: 6px;
                font-size: 13px;
                background-color: #f8f9fa;
            }
            QSpinBox:focus {
                border-color: #3498db;
            }
        """)
        extract_len_layout.addWidget(self.extract_len_spin)
        extract_len_layout.addStretch()
        rules_layout.addLayout(extract_len_layout)
        
        # 文件名最大长度
        max_len_layout = QHBoxLayout()
        max_len_label = QLabel("文件名最大长度:")
        max_len_label.setStyleSheet("font-weight: bold; color: #2c3e50; font-size: 13px;")
        max_len_layout.addWidget(max_len_label)
        
        self.max_len_spin = QSpinBox()
        self.max_len_spin.setRange(20, 100)
        self.max_len_spin.setValue(60)
        self.max_len_spin.setStyleSheet("""
            QSpinBox {
                padding: 8px;
                border: 2px solid #e0e0e0;
                border-radius: 6px;
                font-size: 13px;
                background-color: #f8f9fa;
            }
            QSpinBox:focus {
                border-color: #3498db;
            }
        """)
        max_len_layout.addWidget(self.max_len_spin)
        max_len_layout.addStretch()
        rules_layout.addLayout(max_len_layout)
        
        # 其他选项
        options_layout = QVBoxLayout()
        self.lowercase_checkbox = QCheckBox("转换为小写")
        self.lowercase_checkbox.setChecked(True)
        self.lowercase_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QCheckBox::indicator {
                width: 20px;
                height: 20px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 4px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xIDQuNUw0LjUgOEwxMSAxIiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
            }
        """)
        
        self.space_to_underscore_checkbox = QCheckBox("空格转换为下划线")
        self.space_to_underscore_checkbox.setChecked(True)
        self.space_to_underscore_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QCheckBox::indicator {
                width: 20px;
                height: 20px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 4px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xIDQuNUw0LjUgOEwxMSAxIiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
            }
        """)
        
        options_layout.addWidget(self.lowercase_checkbox)
        options_layout.addWidget(self.space_to_underscore_checkbox)
        rules_layout.addLayout(options_layout)
        
        rules_filter_layout.addWidget(rules_group, 1)  # 设置拉伸比例为1
        
        # 文件类型过滤
        filter_group = QGroupBox("📄 文件类型过滤")
        filter_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #2c3e50;
            }
        """)
        filter_layout = QVBoxLayout(filter_group)
        filter_layout.setSpacing(10)
        
        self.include_images_checkbox = QCheckBox("🖼️ 图片文件 (jpg, png, gif, bmp等)")
        self.include_images_checkbox.setChecked(True)
        self.include_images_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QCheckBox::indicator {
                width: 20px;
                height: 20px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 4px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xIDQuNUw0LjUgOEwxMSAxIiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
            }
        """)
        
        self.include_pdfs_checkbox = QCheckBox("📄 PDF文件")
        self.include_pdfs_checkbox.setChecked(True)
        self.include_pdfs_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QCheckBox::indicator {
                width: 20px;
                height: 20px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 4px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xIDQuNUw0LjUgOEwxMSAxIiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
            }
        """)
        
        self.include_docs_checkbox = QCheckBox("📝 文档文件 (docx, txt等)")
        self.include_docs_checkbox.setChecked(True)
        self.include_docs_checkbox.setStyleSheet("""
            QCheckBox {
                font-size: 13px;
                color: #2c3e50;
                padding: 8px;
            }
            QCheckBox::indicator {
                width: 20px;
                height: 20px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #3498db;
                border-radius: 4px;
                background-color: #3498db;
                image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xIDQuNUw0LjUgOEwxMSAxIiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
            }
        """)
        
        filter_layout.addWidget(self.include_images_checkbox)
        filter_layout.addWidget(self.include_pdfs_checkbox)
        filter_layout.addWidget(self.include_docs_checkbox)
        
        rules_filter_layout.addWidget(filter_group, 1)  # 设置拉伸比例为1，与重命名规则区域平分宽度
        
        # 将水平布局添加到主布局
        layout.addLayout(rules_filter_layout)
        
        # 保存配置按钮
        save_config_btn = QPushButton("💾 保存配置")
        save_config_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 15px 30px;
                border-radius: 8px;
                font-size: 16px;
                font-weight: bold;
                margin: 10px 0;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:pressed {
                background-color: #1e8449;
            }
        """)
        save_config_btn.clicked.connect(self.save_config)
        layout.addWidget(save_config_btn)
        
        layout.addStretch()
        return widget
    
    def create_result_tab(self) -> QWidget:
        """创建结果标签页"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # 结果统计
        stats_group = QGroupBox("执行统计")
        stats_layout = QHBoxLayout(stats_group)
        
        self.total_files_label = QLabel("总文件数: 0")
        self.success_files_label = QLabel("成功: 0")
        self.error_files_label = QLabel("失败: 0")
        
        stats_layout.addWidget(self.total_files_label)
        stats_layout.addWidget(self.success_files_label)
        stats_layout.addWidget(self.error_files_label)
        stats_layout.addStretch()
        
        layout.addWidget(stats_group)
        
        # 详细结果表格
        result_group = QGroupBox("详细结果")
        result_layout = QVBoxLayout(result_group)
        
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(4)
        self.result_table.setHorizontalHeaderLabels(["原文件名", "新文件名", "操作", "时间"])
        self.result_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        
        result_layout.addWidget(self.result_table)
        layout.addWidget(result_group)
        
        # 导出结果按钮
        export_layout = QHBoxLayout()
        self.export_log_btn = QPushButton("导出日志")
        self.export_log_btn.clicked.connect(self.export_log)
        self.export_log_btn.setEnabled(False)
        
        export_layout.addWidget(self.export_log_btn)
        export_layout.addStretch()
        layout.addLayout(export_layout)
        
        layout.addStretch()
        return widget
    
    def select_files_or_folders(self):
        """选择文件或文件夹"""
        # 创建一个自定义对话框，让用户选择模式
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton
        
        dialog = QDialog(self)
        dialog.setWindowTitle("选择模式")
        dialog.setFixedSize(300, 120)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
            }
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
                margin: 5px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QLabel {
                font-size: 14px;
                color: #2c3e50;
                margin: 10px;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        
        # 标题
        title = QLabel("请选择要添加的内容类型：")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        # 按钮布局
        btn_layout = QHBoxLayout()
        
        # 选择文件按钮
        select_files_btn = QPushButton("📄 选择文件")
        select_files_btn.clicked.connect(lambda: self._select_files_and_close(dialog))
        
        # 选择文件夹按钮
        select_folder_btn = QPushButton("📁 选择文件夹")
        select_folder_btn.clicked.connect(lambda: self._select_folder_and_close(dialog))
        
        btn_layout.addWidget(select_files_btn)
        btn_layout.addWidget(select_folder_btn)
        layout.addLayout(btn_layout)
        
        # 显示对话框
        dialog.exec()
    
    def _select_files_and_close(self, dialog):
        """选择文件并关闭对话框"""
        if dialog:
            dialog.accept()
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择要重命名的文件", "",
            "所有文件 (*);;图片文件 (*.jpg *.png *.gif *.bmp);;PDF文件 (*.pdf);;文档文件 (*.docx *.txt)"
        )
        if files:
            self.source_paths.extend([Path(f) for f in files])
            self.update_file_list_display()
    
    def _select_folder_and_close(self, dialog):
        """选择文件夹并关闭对话框"""
        if dialog:
            dialog.accept()
        folder = QFileDialog.getExistingDirectory(self, "选择要重命名的文件夹")
        if folder:
            self.source_paths.append(Path(folder))
            self.update_file_list_display()
    
    def select_files(self):
        """选择文件（保留向后兼容）"""
        self._select_files_and_close(None)
    
    def select_folder(self):
        """选择文件夹（保留向后兼容）"""
        self._select_folder_and_close(None)
    
    def clear_selection(self):
        """清除选择"""
        self.source_paths.clear()
        self.update_file_list_display()
    
    def update_file_list_display(self):
        """更新文件列表显示"""
        if not self.source_paths:
            self.file_list_label.setText("未选择任何文件")
            return
        
        file_count = 0
        for path in self.source_paths:
            if path.is_file():
                file_count += 1
            elif path.is_dir():
                file_count += len(list(path.rglob("*")))
        
        self.file_list_label.setText(f"已选择 {len(self.source_paths)} 个路径，共 {file_count} 个文件")
    
    def select_target_directory(self):
        """选择目标目录"""
        folder = QFileDialog.getExistingDirectory(self, "选择目标目录")
        if folder:
            self.target_dir = Path(folder)
            self.target_path_edit.setText(str(self.target_dir))
    
    def get_config(self) -> Dict[str, Any]:
        """获取当前配置"""
        include_exts = []
        if self.include_images_checkbox.isChecked():
            include_exts.extend(['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tif', 'tiff', 'webp'])
        if self.include_pdfs_checkbox.isChecked():
            include_exts.append('pdf')
        if self.include_docs_checkbox.isChecked():
            include_exts.extend(['docx', 'txt', 'md', 'csv'])
        
        return {
            'extract_len': self.extract_len_spin.value(),
            'max_length': self.max_len_spin.value(),
            'lowercase': self.lowercase_checkbox.isChecked(),
            'space_to_underscore': self.space_to_underscore_checkbox.isChecked(),
            'include_exts': include_exts,
            'copy_mode': self.copy_mode_radio.isChecked()
        }
    
    def start_rename(self):
        """开始重命名"""
        if not self.source_paths:
            QMessageBox.warning(self, "警告", "请先选择要重命名的文件或文件夹")
            return
        
        if not self.target_dir:
            QMessageBox.warning(self, "警告", "请选择目标目录")
            return
        
        # 确认操作
        mode = "复制" if self.copy_mode_radio.isChecked() else "覆盖"
        reply = QMessageBox.question(
            self, "确认操作", 
            f"确定要{mode}重命名 {len(self.source_paths)} 个路径下的文件吗？\n"
            f"目标目录: {self.target_dir}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # 禁用界面
        self.execute_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # 清空日志
        self.log_text.clear()
        self.log_text.append(f"开始执行重命名操作...\n模式: {mode}\n目标目录: {self.target_dir}")
        
        # 创建工作线程
        copy_mode = self.copy_mode_radio.isChecked()
        config = self.get_config()
        
        self.worker = RenameWorker(self.source_paths, self.target_dir, config, copy_mode, self.log_message)
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.file_processed.connect(self.log_file_processed)
        self.worker.finished.connect(self.rename_finished)
        self.worker.error_occurred.connect(self.log_error)
        
        self.worker.start()
    
    def update_progress(self, current: int, total: int):
        """更新进度条"""
        self.progress_bar.setValue(int(current / total * 100))
        self.statusBar().showMessage(f"处理进度: {current}/{total}")
    
    def log_file_processed(self, old_name: str, new_name: str):
        """记录文件处理日志"""
        self.log_text.append(f"✓ {old_name} → {new_name}")
    
    def log_error(self, error_msg: str):
        """记录错误日志"""
        self.log_text.append(f"✗ 错误: {error_msg}")
    
    def rename_finished(self, results: Dict[str, Any]):
        """重命名完成"""
        self.execute_btn.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.rollback_btn.setEnabled(True)
        self.export_log_btn.setEnabled(True)
        
        # 更新统计信息
        self.total_files_label.setText(f"总文件数: {results['total']}")
        self.success_files_label.setText(f"成功: {results['success']}")
        self.error_files_label.setText(f"失败: {results['error']}")
        
        # 更新结果表格
        self.update_result_table(results['rename_log'])
        
        # 保存重命名日志
        self.rename_log = results['rename_log']
        
        # 显示完成消息
        QMessageBox.information(
            self, "完成", 
            f"重命名操作完成！\n"
            f"成功: {results['success']} 个文件\n"
            f"失败: {results['error']} 个文件"
        )
        
        self.statusBar().showMessage("重命名操作完成")
    
    def update_result_table(self, rename_log: List[Dict]):
        """更新结果表格"""
        self.result_table.setRowCount(len(rename_log))
        
        for row, log_entry in enumerate(rename_log):
            old_name = Path(log_entry['old_path']).name
            new_name = Path(log_entry['new_path']).name
            action = log_entry['action']
            timestamp = log_entry['timestamp']
            
            self.result_table.setItem(row, 0, QTableWidgetItem(old_name))
            self.result_table.setItem(row, 1, QTableWidgetItem(new_name))
            self.result_table.setItem(row, 2, QTableWidgetItem(action))
            self.result_table.setItem(row, 3, QTableWidgetItem(timestamp))
    
    def rollback_operations(self):
        """回滚操作"""
        if not self.rename_log:
            QMessageBox.information(self, "提示", "没有可回滚的操作")
            return
        
        reply = QMessageBox.question(
            self, "确认回滚", 
            f"确定要回滚 {len(self.rename_log)} 个重命名操作吗？\n"
            "这将尝试恢复所有文件到重命名前的状态。",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # 执行回滚
        success_count = 0
        error_count = 0
        
        for log_entry in reversed(self.rename_log):  # 倒序回滚
            try:
                old_path = Path(log_entry['old_path'])
                new_path = Path(log_entry['new_path'])
                
                if log_entry['action'] == 'copied':
                    # 复制模式：删除新文件
                    if new_path.exists():
                        new_path.unlink()
                else:
                    # 覆盖模式：恢复原文件名
                    if new_path.exists():
                        new_path.rename(old_path)
                
                success_count += 1
                
            except Exception as e:
                error_count += 1
                self.log_text.append(f"回滚失败: {log_entry['old_path']} - {str(e)}")
        
        # 显示回滚结果
        QMessageBox.information(
            self, "回滚完成", 
            f"回滚操作完成！\n"
            f"成功: {success_count} 个文件\n"
            f"失败: {error_count} 个文件"
        )
        
        # 清空日志和表格
        self.rename_log.clear()
        self.result_table.setRowCount(0)
        self.rollback_btn.setEnabled(False)
        self.export_log_btn.setEnabled(False)
        
        self.statusBar().showMessage("回滚操作完成")
    
    def export_log(self):
        """导出日志"""
        if not self.rename_log:
            QMessageBox.information(self, "提示", "没有可导出的日志")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出日志", f"rename_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "JSON文件 (*.json)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(self.rename_log, f, ensure_ascii=False, indent=2)
                QMessageBox.information(self, "成功", f"日志已导出到: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")
    
    def load_config(self):
        """加载配置"""
        try:
            # 加载API配置
            config_path = Path(__file__).parent / "config.json"
            if config_path.exists():
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    api_key = config.get('deepseek_api_key', '')
                    if api_key and api_key != "your_api_key_here":
                        self.api_key_input.setText(api_key)
                        self.api_status_label.setText("API状态: 已配置")
                        self.api_status_label.setStyleSheet("color: green; margin: 5px;")
                    else:
                        self.api_status_label.setText("API状态: 未配置")
                        self.api_status_label.setStyleSheet("color: orange; margin: 5px;")
        except Exception as e:
            print(f"加载配置失败: {e}")
    
    def save_config(self):
        """保存所有配置（包括API配置和重命名规则配置）"""
        try:
            # 保存API配置
            api_key = self.api_key_input.text().strip()
            if api_key and api_key != "your_api_key_here":
                api_config = {
                    "deepseek_api_key": api_key,
                    "description": "请将your_api_key_here替换为你的DeepSeek API密钥"
                }
                
                api_config_path = Path(__file__).parent / "config.json"
                with open(api_config_path, 'w', encoding='utf-8') as f:
                    json.dump(api_config, f, indent=4, ensure_ascii=False)
                
                # 更新API状态显示
                self.api_status_label.setText("API状态: 已配置")
                self.api_status_label.setStyleSheet("color: green; margin: 5px;")
                
                # 重新加载 DeepSeek API 密钥
                try:
                    from deepseek_api_service import deepseek_service
                    deepseek_service.reload_api_key()
                except ImportError:
                    pass
            
            # 保存重命名规则配置
            app_config = {
                "extract_len": self.extract_len_spin.value(),
                "max_len": self.max_len_spin.value(),
                "lowercase": self.lowercase_checkbox.isChecked(),
                "space_to_underscore": self.space_to_underscore_checkbox.isChecked(),
                "include_images": self.include_images_checkbox.isChecked(),
                "include_pdfs": self.include_pdfs_checkbox.isChecked(),
                "include_docs": self.include_docs_checkbox.isChecked(),
                "copy_mode": self.copy_mode_radio.isChecked()
            }
            
            app_config_path = Path(__file__).parent / "app_config.json"
            with open(app_config_path, 'w', encoding='utf-8') as f:
                json.dump(app_config, f, indent=4, ensure_ascii=False)
            
            QMessageBox.information(self, "成功", "所有配置已保存！")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存配置失败: {e}")
    

    
    def test_api_key(self):
        """测试API密钥"""
        try:
            api_key = self.api_key_input.text().strip()
            
            if not api_key or api_key == "your_api_key_here":
                QMessageBox.warning(self, "警告", "请先配置有效的API密钥")
                return
            
            # 测试API连接
            import requests
            
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "deepseek-chat",
                "messages": [{"role": "user", "content": "Hello"}],
                "max_tokens": 10
            }
            
            response = requests.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=10
            )
            
            if response.status_code == 200:
                QMessageBox.information(self, "成功", "API密钥配置正确！")
                self.api_status_label.setText("API状态: 测试成功")
                self.api_status_label.setStyleSheet("color: green; margin: 5px;")
            else:
                QMessageBox.warning(self, "警告", f"API密钥测试失败: {response.status_code}")
                self.api_status_label.setText("API状态: 测试失败")
                self.api_status_label.setStyleSheet("color: red; margin: 5px;")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"测试API密钥失败: {e}")
            self.api_status_label.setText("API状态: 测试失败")
            self.api_status_label.setStyleSheet("color: red; margin: 5px;")
    
    def clear_api_key(self):
        """清空API密钥"""
        self.api_key_input.clear()
        self.api_status_label.setText("API状态: 未配置")
        self.api_status_label.setStyleSheet("color: gray; margin: 5px;")
    
    def toggle_password_visibility(self):
        """切换API密钥显示/隐藏"""
        if self.api_key_input.echoMode() == QLineEdit.EchoMode.Password:
            # 当前是密码模式，切换到明文模式
            self.api_key_input.setEchoMode(QLineEdit.EchoMode.Normal)
            self.toggle_password_btn.setText("🙈")
            self.toggle_password_btn.setToolTip("隐藏API密钥")
        else:
            # 当前是明文模式，切换到密码模式
            self.api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
            self.toggle_password_btn.setText("👁")
            self.toggle_password_btn.setToolTip("显示API密钥")


def main():
    """主函数"""
    app = QApplication(sys.argv)
    
    # 设置应用样式
    app.setStyle('Fusion')
    
    # 创建主窗口
    window = FileRenamerGUI()
    window.show()
    
    # 运行应用
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
